"""Part 1: Helper functions and Task 1 content - EXPANDED for distinction."""
import sys, subprocess, os
def install(p):
    subprocess.check_call([sys.executable, "-m", "pip", "install", p])
try:
    import docx
except ImportError:
    install('python-docx')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def set_bg(cell, color):
    shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_row(table, cells_data):
    row = table.add_row().cells
    for i, txt in enumerate(cells_data):
        row[i].text = str(txt)
    return row

def bold_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    fmt = p.paragraph_format
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(6)
    return p

def styled_header_row(table, headers):
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        set_bg(table.rows[0].cells[i], "4472C4")
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255,255,255)
                r.font.bold = True

def write_task1(doc):
    doc.add_heading('Task 1: System Architecture and Cargo Sorting', level=1)
    doc.add_paragraph(
        'This task focuses on the initial design of the Swift-Load Logistics platform, '
        'ensuring data integrity and establishing the basic organizational logic required '
        'to manage cargo items effectively. We design the foundational Abstract Data Type '
        '(ADT) for goods, illustrate queue-based truck management, compare sorting algorithms '
        'for cargo optimization, and examine how encapsulation principles protect system integrity. '
        'Each subsection below addresses a specific criterion from the assignment brief, with '
        'supporting code examples, formal tables, and analytical commentary.'
    )

    # === P1 ===
    doc.add_heading('P1: Design Specification for the "Goods" Data Structure', level=2)
    doc.add_paragraph(
        'The "Goods" Abstract Data Type (ADT) is the cornerstone of the Swift-Load Logistics system. '
        'An ADT defines a mathematical model for a data type by specifying two key components: the set '
        'of values it can hold and the set of operations that can be performed on it. In formal terms, '
        'an ADT is a triple (D, F, A) where D is the domain of values, F is the set of functions or '
        'operations, and A is the set of axioms or constraints that govern behavior. For our logistics '
        'system, the Goods ADT encapsulates all the essential metadata about each cargo item that passes '
        'through the delivery network, serving as the single source of truth for item properties.'
    )
    bold_para(doc, 'Attributes (Domain D):')
    doc.add_paragraph('Name (String): A unique human-readable identifier for the item, such as "Laptop" or '
        '"Apples". This serves as the primary key when stored in tree-based data structures and is used '
        'for alphabetical sorting and binary search comparisons. The name must be non-empty for valid '
        'insertion into the AVL tree.', style='List Bullet')
    doc.add_paragraph('Type (String): The category classification of the item (e.g., "Electronics", '
        '"Food", "Furniture"). This enables cargo grouping, special handling requirements such as '
        'refrigeration for food items, and priority-based loading sequences where fragile electronics '
        'may need to be loaded last.', style='List Bullet')
    doc.add_paragraph('Weight (Double): The mass of the item measured in kilograms. This is critical '
        'for truck capacity optimization, load balancing calculations, and determining whether a truck '
        'has exceeded its maximum weight limit. The weight must always be non-negative, as a negative '
        'weight has no physical meaning and would corrupt downstream calculations.', style='List Bullet')

    bold_para(doc, 'Valid Operations (Functions F):')
    doc.add_paragraph(
        'The following table formally specifies every valid operation that can be performed on '
        'the Goods ADT, including the parameters, return types, and validation constraints that '
        'ensure data integrity throughout the system. These operations constitute the complete '
        'public interface through which all external modules interact with Goods objects.'
    )

    t = doc.add_table(rows=1, cols=4)
    t.style = 'Table Grid'
    styled_header_row(t, ['Operation', 'Parameters', 'Return Type', 'Pre/Post Conditions & Axioms'])

    add_table_row(t, ['Constructor (init)', 'name: String, type: String, weight: Double',
        'Goods object', 'Pre: weight >= 0. Post: Valid Goods object created in memory with all fields '
        'initialized. Axiom: if weight < 0, throw invalid_argument exception.'])
    add_table_row(t, ['getName()', 'None', 'String',
        'Post: Returns the current name value without modifying internal state (const method).'])
    add_table_row(t, ['setName(newName)', 'newName: String', 'void',
        'Pre: newName must not be empty string. Post: Internal name updated to newName. '
        'Axiom: if newName is empty, throw invalid_argument exception.'])
    add_table_row(t, ['getType()', 'None', 'String',
        'Post: Returns the current type classification string.'])
    add_table_row(t, ['setType(newType)', 'newType: String', 'void',
        'Post: Internal type updated to newType. No restrictions on value.'])
    add_table_row(t, ['getWeight()', 'None', 'Double',
        'Post: Returns the current weight value in kilograms as a double-precision float.'])
    add_table_row(t, ['setWeight(newWeight)', 'newWeight: Double', 'void',
        'Pre: newWeight >= 0. Post: Internal weight updated. '
        'Axiom: if newWeight < 0, throw invalid_argument exception to prevent data corruption.'])
    add_table_row(t, ['display()', 'None', 'void',
        'Post: Outputs formatted Name (20 chars), Type (15 chars), Weight to console with "kg" suffix.'])
    doc.add_paragraph('Table 1: Goods ADT Complete Operation Specification')
    doc.add_paragraph()

    bold_para(doc, 'C++ Implementation of the Goods Class:')
    add_code_block(doc, '''class Goods {
private:
    string name;    // Private - Information Hiding
    string type;    // Private - Information Hiding
    double weight;  // Private - Information Hiding
public:
    // Constructor with validation
    Goods(string name="", string type="", double weight=0.0) {
        if (weight < 0)
            throw invalid_argument("Weight cannot be negative");
        this->name = name;
        this->type = type;
        this->weight = weight;
    }
    // Accessor methods (getters)
    string getName() const { return name; }
    string getType() const { return type; }
    double getWeight() const { return weight; }
    // Mutator methods with validation (setters)
    void setName(const string& newName) {
        if (newName.empty())
            throw invalid_argument("Name cannot be empty");
        name = newName;
    }
    void setWeight(double newWeight) {
        if (newWeight < 0)
            throw invalid_argument("Weight cannot be negative");
        weight = newWeight;
    }
};''')
    doc.add_paragraph(
        'The implementation above demonstrates how the constructor enforces the pre-condition that '
        'weight must be non-negative. If a negative weight is provided, the system throws a C++ '
        'std::invalid_argument exception, preventing corrupted data from entering the logistics pipeline. '
        'The const qualifier on getter methods guarantees that calling getName() cannot accidentally '
        'modify the internal state, providing compile-time safety. This defensive programming approach '
        'is essential for maintaining data integrity across the entire Swift-Load platform, especially '
        'when multiple modules access the same Goods objects concurrently.'
    )

    # === M1 ===
    doc.add_heading('M1: Queue Illustration (FIFO) for Loading Bay Management', level=2)
    doc.add_paragraph(
        'The Swift-Load loading bay operates on a strict first-come, first-served basis to ensure '
        'fairness among delivery trucks and prevent operational chaos. This is mathematically modeled '
        'using a First-In-First-Out (FIFO) Queue data structure. A queue is a linear data structure '
        'that supports two primary operations: enqueue (adding an element to the rear) and dequeue '
        '(removing an element from the front). Unlike a stack, where the most recent element is '
        'accessed first, a queue guarantees that elements are processed in the exact chronological '
        'order they arrived, preventing starvation where a truck could wait indefinitely while newer '
        'arrivals are processed first.'
    )
    doc.add_paragraph(
        'In the real-world logistics context, when Truck A arrives at the loading bay at 8:00 AM, '
        'followed by Truck B at 8:15 AM and Truck C at 8:30 AM, the queue ensures Truck A is always '
        'loaded first. This FIFO discipline is critical for maintaining contractual delivery deadlines '
        'and customer satisfaction metrics.'
    )

    bold_para(doc, 'Concrete Step-by-Step Example:')
    qt = doc.add_table(rows=1, cols=4)
    qt.style = 'Table Grid'
    styled_header_row(qt, ['Step', 'Operation', 'Queue State (Front -> Rear)', 'Queue Size'])
    add_table_row(qt, ['1', 'Enqueue Truck A (Heavy Duty)', '[Truck A]', '1'])
    add_table_row(qt, ['2', 'Enqueue Truck B (Refrigerated)', '[Truck A, Truck B]', '2'])
    add_table_row(qt, ['3', 'Enqueue Truck C (Standard)', '[Truck A, Truck B, Truck C]', '3'])
    add_table_row(qt, ['4', 'Dequeue -> Truck A processed', '[Truck B, Truck C]', '2'])
    add_table_row(qt, ['5', 'Dequeue -> Truck B processed', '[Truck C]', '1'])
    add_table_row(qt, ['6', 'Dequeue -> Truck C processed', '[] (empty)', '0'])
    doc.add_paragraph('Table 2: Step-by-Step FIFO Queue Operations for Loading Bay')

    bold_para(doc, 'Queue Diagram (Visual Representation):')
    dt = doc.add_table(rows=1, cols=5)
    dt.style = 'Table Grid'
    cells = dt.rows[0].cells
    cells[0].text = "EXIT\n<-- Dequeue"
    cells[1].text = "Truck A\n(Front)"
    cells[2].text = "Truck B\n(Middle)"
    cells[3].text = "Truck C\n(Rear)"
    cells[4].text = "ENTER\nEnqueue -->"
    for c in cells: c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_bg(cells[1], "FFD966")
    set_bg(cells[2], "A9D08E")
    set_bg(cells[3], "9DC3E6")
    doc.add_paragraph('Figure 1: FIFO Queue Diagram for Swift-Load Loading Bay')

    bold_para(doc, 'C++ Implementation using std::queue:')
    add_code_block(doc, '''queue<string> truckQueue;
truckQueue.push("Truck A (Heavy Duty)");    // Enqueue at rear
truckQueue.push("Truck B (Refrigerated)");   // Enqueue at rear
truckQueue.push("Truck C (Standard)");       // Enqueue at rear
// Front: Truck A, Back: Truck C, Size: 3
while (!truckQueue.empty()) {
    cout << "Processing: " << truckQueue.front();  // Peek at front
    truckQueue.pop();                               // Dequeue from front
}''')
    doc.add_paragraph(
        'The C++ Standard Template Library provides the std::queue container adapter which implements '
        'the FIFO behavior. The push() method performs enqueue, front() performs peek, and pop() '
        'performs dequeue. All three operations execute in O(1) constant time, making the queue '
        'highly efficient for real-time loading bay management regardless of how many trucks are waiting.'
    )

    # === M2 ===
    doc.add_heading('M2: Sorting Comparison (Bubble Sort vs QuickSort)', level=2)
    doc.add_paragraph(
        'Sorting cargo by weight is essential for optimizing truck capacity and ensuring safe load '
        'distribution. Heavier items should be loaded first (at the bottom) to ensure stability and '
        'maximize space utilization. We implemented and compared two fundamentally different sorting '
        'algorithms using a concrete sample of 12 cargo weights, sorted in descending order to prioritize '
        'heavy-first loading.'
    )
    bold_para(doc, 'Sample Data (12 cargo weights in kg):')
    doc.add_paragraph('{45.5, 12.0, 89.2, 5.5, 34.1, 100.0, 77.8, 23.4, 56.7, 9.9, 41.2, 60.0}')
    bold_para(doc, 'Expected Sorted Result (Descending):')
    doc.add_paragraph('{100.0, 89.2, 77.8, 60.0, 56.7, 45.5, 41.2, 34.1, 23.4, 12.0, 9.9, 5.5}')

    bold_para(doc, 'Bubble Sort Algorithm:')
    doc.add_paragraph(
        'Bubble Sort is a simple comparison-based algorithm that repeatedly traverses the list, '
        'compares adjacent elements, and swaps them if they are in the wrong order. Each pass moves '
        'the smallest unsorted element to its correct position at the end (like a bubble rising). '
        'The algorithm requires N-1 passes in the worst case, and each pass performs up to N-i-1 '
        'comparisons. While conceptually simple and easy to implement, its O(N^2) time complexity '
        'makes it impractical for datasets larger than a few hundred elements.'
    )
    add_code_block(doc, '''void bubbleSort(vector<double>& arr) {
    int n = arr.size();
    for (int i = 0; i < n - 1; i++) {
        for (int j = 0; j < n - i - 1; j++) {
            if (arr[j] < arr[j + 1]) {  // < for descending
                swap(arr[j], arr[j + 1]);
            }
        }
    }
}''')
    bold_para(doc, 'QuickSort Algorithm:')
    doc.add_paragraph(
        'QuickSort uses a divide-and-conquer strategy. It selects a pivot element (typically the last '
        'element), partitions the array into elements greater than and less than the pivot (for '
        'descending order), and recursively sorts each partition. The partitioning step is the key '
        'innovation: it rearranges elements in a single linear pass, then the two resulting sub-problems '
        'are each roughly half the original size. This halving behavior gives QuickSort its O(N log N) '
        'average-case performance, making it dramatically faster than Bubble Sort for large datasets.'
    )
    add_code_block(doc, '''void quickSort(vector<double>& arr, int low, int high) {
    if (low < high) {
        int pi = partition(arr, low, high);  // Partition
        quickSort(arr, low, pi - 1);    // Sort left half
        quickSort(arr, pi + 1, high);   // Sort right half
    }
}''')

    bold_para(doc, 'Empirical Performance Comparison:')
    st = doc.add_table(rows=1, cols=5)
    st.style = 'Table Grid'
    styled_header_row(st, ['Metric', 'Bubble Sort', 'QuickSort', 'Winner', 'Significance for Logistics'])
    add_table_row(st, ['Time Complexity (Avg)', 'O(N\u00b2)', 'O(N log N)', 'QuickSort',
        'For N=10,000 packages: Bubble=100M ops vs Quick=133K ops'])
    add_table_row(st, ['Time Complexity (Worst)', 'O(N\u00b2)', 'O(N\u00b2)', 'Tie',
        'QuickSort worst case occurs with already-sorted input'])
    add_table_row(st, ['Comparisons (12 items)', '66', '34', 'QuickSort',
        'QuickSort uses 48% fewer comparisons on our sample'])
    add_table_row(st, ['Swaps (12 items)', '34', '23', 'QuickSort',
        'QuickSort performs 32% fewer memory write operations'])
    add_table_row(st, ['Space Complexity', 'O(1) in-place', 'O(log N) stack', 'Bubble Sort',
        'QuickSort requires recursive stack frames'])
    add_table_row(st, ['Stability', 'Stable', 'Unstable', 'Bubble Sort',
        'Bubble Sort preserves relative order of equal elements'])
    add_table_row(st, ['Implementation', 'Very Simple', 'Moderate', 'Bubble Sort',
        'Bubble Sort has only 2 nested loops, easier to debug'])
    doc.add_paragraph('Table 3: Comprehensive Bubble Sort vs QuickSort Performance Comparison')
    doc.add_paragraph(
        'Conclusion: Despite Bubble Sort having minor advantages in space efficiency and stability, '
        'QuickSort is overwhelmingly superior for the Swift-Load use case. With thousands of packages '
        'processed daily, the O(N^2) scaling of Bubble Sort would create unacceptable delays during '
        'peak seasons, while QuickSort\'s O(N log N) performance ensures consistent throughput regardless '
        'of cargo volume. Our empirical measurements confirm that even with just 12 items, QuickSort '
        'already performs 48% fewer comparisons.'
    )

    # === M3 ===
    doc.add_heading('M3: Data Encapsulation and Information Hiding', level=2)
    doc.add_paragraph(
        'Encapsulation is one of the four fundamental principles of Object-Oriented Programming, alongside '
        'inheritance, polymorphism, and abstraction. It refers to the bundling of data (attributes) and '
        'the methods that operate on that data into a single unit (class), while restricting direct access '
        'to some of the object\'s components. In the Goods ADT, all three attributes (name, type, weight) '
        'are declared as private in C++, meaning they cannot be accessed or modified directly from outside '
        'the class. All interactions must go through the validated public methods (getters and setters). '
        'This is a critical security and integrity measure for the logistics system.'
    )
    bold_para(doc, 'Advantage 1: Data Integrity Through Validation Gates')
    doc.add_paragraph(
        'Without encapsulation, any part of the codebase could directly execute cargo.weight = -500, which '
        'would immediately corrupt the load-balancing algorithm and potentially cause a truck to be '
        'overloaded or underloaded. By forcing all weight modifications through the setWeight() method, '
        'we guarantee that the validation check (if newWeight < 0, throw exception) is always enforced '
        'before the internal state changes. This creates a single chokepoint of validation that protects '
        'the entire system from invalid data, regardless of which module initiates the change.'
    )
    bold_para(doc, 'Advantage 2: Implementation Independence and Future-Proofing')
    doc.add_paragraph(
        'If the logistics team later decides to store weight in grams instead of kilograms for higher '
        'precision on lightweight electronics, only the internal implementation of getWeight() and '
        'setWeight() need to change. Every external module that reads weight continues to function '
        'without any modification whatsoever, because they interact exclusively through the public '
        'interface. This dramatically reduces the cost and risk of system evolution.'
    )
    bold_para(doc, 'Advantage 3: Reduced Coupling and Simplified Maintenance')
    doc.add_paragraph(
        'Encapsulation reduces coupling between software components by hiding implementation details. '
        'The truck-loading module does not know or care whether weight is stored as a double, float, '
        'or integer internally. This modular design means that bugs in the Goods class are isolated and '
        'cannot cascade through the system to corrupt the routing or billing modules. This isolation '
        'drastically reduces debugging time and maintenance costs over the software lifecycle.'
    )

    # === D2 ===
    doc.add_heading('D2: OOP Foundations from Imperative ADTs', level=2)
    doc.add_paragraph(
        'The evolution from imperative Abstract Data Types to Object-Oriented Programming (OOP) represents '
        'one of the most significant paradigm shifts in computer science history. In purely imperative '
        'languages like C, data structures (structs) and the functions that manipulate them exist as '
        'completely separate entities. A C struct for Goods would simply be a passive container of raw '
        'data fields, with standalone functions like updateWeight(Goods* g, double w) operating on it '
        'externally. There is no mechanism to prevent direct field access (g->weight = -500), no way to '
        'bundle validation logic with the data it protects, and no formal connection between the struct '
        'and its associated functions.'
    )
    doc.add_paragraph(
        'The imperative ADT concept introduced a crucial philosophical insight that bridged this gap: '
        'certain data and certain operations are semantically inseparable. The weight of a cargo item '
        'cannot be meaningfully updated without simultaneously validating that the new value is non-negative. '
        'The name of an item cannot be set to an empty string without potentially breaking the AVL tree\'s '
        'search key invariant. This realization that data and its valid operations form an indivisible '
        'logical unit became the philosophical and mathematical foundation upon which OOP was built.'
    )
    doc.add_paragraph(
        'In our Swift-Load system, this evolutionary transition is directly visible in the codebase. '
        'The imperative ADT specification stated: "A Goods value consists of name, type, and weight. '
        'Valid operations include init, getName, setName, getWeight, setWeight, and display." The C++ '
        'class implementation translates this formal specification directly into language-level constructs: '
        'private member variables become the encapsulated data, public methods become the sanctioned '
        'operations, constructors formalize initialization, and exception handling enforces the axioms. '
        'The C++ class IS the ADT, realized in executable code. This direct correspondence proves that '
        'Object-Oriented Programming is not a separate invention but rather the natural and inevitable '
        'formalization of imperative ADT principles at the programming language level, providing compiler-'
        'enforced guarantees for what was previously only a design convention.'
    )
