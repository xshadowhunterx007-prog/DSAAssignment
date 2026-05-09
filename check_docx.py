import docx
import sys

def check_docx_for_old_code(filepath):
    try:
        doc = docx.Document(filepath)
    except Exception as e:
        print(f"Error opening document: {e}")
        return

    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
                
    content = "\n".join(full_text)
    
    old_snippets = [
        ("Old Constructor (no name check)", 'if (weight < 0) {'),
        ("Old BubbleSort (int n)", 'int n = arr.size();'),
        ("Old InsertNode (return node)", 'return node;'),
        ("Old Test 2 (no exception check)", 'Goods nullGoods("", "Misc", 10.0);'),
    ]
    
    found_any = False
    for name, snippet in old_snippets:
        if snippet in content:
            print(f"[FOUND] {name}: '{snippet}' was found in the document.")
            found_any = True
        else:
            print(f"[CLEAN] {name}: Not found.")
            
    if not found_any:
        print("\nGood news! None of the specific old code snippets were found in the text.")
    else:
        print("\nYou will need to update the found snippets in your Word document to match the new main.cpp.")

if __name__ == "__main__":
    filepath = r"d:\school stuff\Final Assignment 6th semester\Data Structure\(Ai fixed) SBS24010097_PhoneMyatMin_DataStructuresAndAlgorithms_FinalAssignment.docx"
    check_docx_for_old_code(filepath)
