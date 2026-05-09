import docx
import sys

def analyze_docx_structure(filepath):
    try:
        doc = docx.Document(filepath)
    except Exception as e:
        print(f"Error opening document: {e}")
        return

    with open("d:\\school stuff\\Final Assignment 6th semester\\Data Structure\\docx_structure.txt", "w", encoding="utf-8") as f:
        f.write("--- HEADINGS IN THE DOCUMENT ---\n")
        heading_count = 0
        for i, para in enumerate(doc.paragraphs):
            if para.style.name.startswith('Heading'):
                f.write(f"[{i}] {para.style.name}: {para.text[:100]}\n")
                heading_count += 1
                
        f.write(f"\nTotal Headings Found: {heading_count}\n")
        
        f.write("\n--- TABLES IN THE DOCUMENT ---\n")
        for i, table in enumerate(doc.tables):
            try:
                first_row_text = [cell.text.replace('\\n', ' ')[:30] for cell in table.rows[0].cells]
                f.write(f"Table {i}: Columns={len(table.columns)}, Headers={first_row_text}\n")
            except:
                f.write(f"Table {i}: (Empty or unstructured)\n")

        f.write("\n--- FINDING CODE SNIPPETS ---\n")
        for i, para in enumerate(doc.paragraphs):
            if 'bubbleSort' in para.text:
                f.write(f"Found 'bubbleSort' at paragraph {i}: {para.text[:50]}...\n")
            if 'class Goods {' in para.text or 'class Goods{' in para.text:
                f.write(f"Found 'class Goods' at paragraph {i}: {para.text[:50]}...\n")
            if 'Goods(string name' in para.text:
                f.write(f"Found Goods constructor at paragraph {i}: {para.text[:50]}...\n")



if __name__ == "__main__":
    filepath = r"d:\school stuff\Final Assignment 6th semester\Data Structure\(Ai fixed) SBS24010097_PhoneMyatMin_DataStructuresAndAlgorithms_FinalAssignment.docx"
    analyze_docx_structure(filepath)
